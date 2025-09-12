# docx_parse.py
# ------------------------------------------------------------
# Мақсаты:
#   .docx ішінен барлық мәтінді жолдарға бөліп шығару.
#   Word-тың "автоматты" тізім нөмірлерін (1., a), i.) қайта құру.
#   Буллеттерді '-' ретінде көрсету.
#
# Қайтарады:
#   list[str]  → әр жол бір логикалық жол
#
# Тәуелділік:
#   pip install python-docx
# ------------------------------------------------------------

from __future__ import annotations
from typing import List, Tuple, Optional, Dict
from docx import Document

# WordprocessingML namespace
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# ----------------- Көмекші түрлендірулер -----------------

def _int_or_none(v) -> Optional[int]:
    try:
        return int(v)
    except Exception:
        return None

def _to_roman(n: int, upper: bool = False) -> str:
    # Қарапайым араб→рим түрлендіргіш (1..3999 жеткілікті)
    vals = [
        (1000, "M"), (900, "CM"),
        (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"),
        (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"),
        (5, "V"), (4, "IV"),
        (1, "I"),
    ]
    res = []
    num = max(1, min(3999, n))
    for v, s in vals:
        while num >= v:
            res.append(s)
            num -= v
    s = "".join(res)
    return s if upper else s.lower()

def _to_letter(n: int, upper: bool = False) -> str:
    # 1->a, 2->b ... 27->aa т.б.
    n = max(1, n)
    result = []
    while n > 0:
        n, rem = divmod(n - 1, 26)
        result.append(chr(ord('A') + rem) if upper else chr(ord('a') + rem))
    return "".join(reversed(result))

def _normalize_space(s: str) -> str:
    import re
    return re.sub(r"\s+", " ", s).strip()

# ----------------- Numbering форматін табу -----------------

class NumberingResolver:
    """
    Word нөмірлеуін анықтау:
      numId + ilvl → numFmt ('decimal', 'lowerLetter', 'upperLetter', 'bullet', 'lowerRoman', т.б.)
    """
    def __init__(self, doc: Document):
        self.doc = doc
        self._cache_fmt: Dict[Tuple[str, int], str] = {}
        # numId -> abstractNumId кеші
        self._cache_abs: Dict[str, str] = {}

        # numbering_part кей құжаттарда болмауы мүмкін (мысалы, нөмірлеу мүлдем жоқ болса)
        self._numbering = None
        try:
            self._numbering = self.doc.part.numbering_part.element
        except Exception:
            self._numbering = None

    def _get_abstract_num_id(self, num_id: str) -> Optional[str]:
        if self._numbering is None:
            return None
        if num_id in self._cache_abs:
            return self._cache_abs[num_id]
        # //w:num[@w:numId='X']/w:abstractNumId/@w:val
        path = f".//w:num[@w:numId='{num_id}']/w:abstractNumId"
        nodes = self._numbering.xpath(path, namespaces=NS)
        if nodes:
            abs_id = nodes[0].get(f"{{{NS['w']}}}val")
            self._cache_abs[num_id] = abs_id
            return abs_id
        return None

    def get_num_fmt(self, num_id: str, ilvl: int) -> Optional[str]:
        """
        numId+ilvl → 'decimal' | 'lowerLetter' | 'upperLetter' | 'lowerRoman' | 'upperRoman' | 'bullet' | ...
        Болмаса None.
        """
        key = (num_id, ilvl)
        if key in self._cache_fmt:
            return self._cache_fmt[key]
        abs_id = self._get_abstract_num_id(num_id)
        if not abs_id or self._numbering is None:
            return None
        # //w:abstractNum[@w:abstractNumId='Y']/w:lvl[@w:ilvl='Z']/w:numFmt/@w:val
        path = f".//w:abstractNum[@w:abstractNumId='{abs_id}']/w:lvl[@w:ilvl='{ilvl}']/w:numFmt"
        nodes = self._numbering.findall(path, namespaces=NS)

        if nodes:
            fmt = nodes[0].get(f"{{{NS['w']}}}val")
            self._cache_fmt[key] = fmt
            return fmt
        return None

# ----------------- Автонөмірлеуді қайта құру -----------------

class ListMarkerBuilder:
    """
    Бір numId үшін деңгейлік санауыш стегі:
      level 0 → 1., 2., 3.  (немесе fmt-ке сай)
      level 1 → a), b), c)
      level 2 → i., ii., iii.
    """
    def __init__(self, resolver: NumberingResolver):
        self.resolver = resolver
        # numId → stack[list[int]] (әр деңгейдің ағымдағы санауышы)
        self.stacks: Dict[str, List[int]] = {}

    def next_marker(self, num_id: str, ilvl: int) -> str:
        stack = self.stacks.get(num_id, [])
        # Стекті ағымдағы деңгейге дейін қысқарту/созу
        if len(stack) > ilvl + 1:
            stack = stack[: ilvl + 1]
        while len(stack) < ilvl + 1:
            stack.append(0)
        # Ағымдағы деңгейді инкременттеу
        stack[ilvl] += 1
        # Төменгі деңгейлерді нөлдеу (қауіпсіздік)
        for i in range(ilvl + 1, len(stack)):
            stack[i] = 0
        self.stacks[num_id] = stack

        # Форматты анықтау
        fmt = self.resolver.get_num_fmt(num_id, ilvl)
        n = stack[ilvl]

        # fmt болмай қалса — келісілген дефолт:
        # 0-деңгей: 1., 1) ; 1-деңгей: a) ; 2-деңгей: i.
        if fmt is None:
            if ilvl == 0:
                return f"{n}."
            elif ilvl == 1:
                return f"{_to_letter(n)})"
            elif ilvl == 2:
                return f"{_to_roman(n)}."
            else:
                # терең болса — жай буллет
                return "-"

        fmt = fmt.lower()
        if fmt == "decimal":
            return f"{n}."
        if fmt == "lowerletter":
            return f"{_to_letter(n)})"
        if fmt == "upperletter":
            return f"{_to_letter(n, upper=True)})"
        if fmt == "lowerroman":
            return f"{_to_roman(n)}."
        if fmt == "upperroman":
            return f"{_to_roman(n, upper=True)}."
        if fmt == "bullet":
            return "-"
        # Басқа форматтар — жалпы белгі
        return f"{n}."

# ----------------- Негізгі экстрактор -----------------

def extract_text(docx_path: str) -> List[str]:
    """
    .docx → list[str]
    - Параграфтар мен кестелерден мәтін жинайды.
    - Word-тың автоматты нөмірлеуін қайта құрады.
    - Бос жолдарды алып тастайды.
    """
    doc = Document(docx_path)
    resolver = NumberingResolver(doc)
    marker_builder = ListMarkerBuilder(resolver)

    lines: List[str] = []

    def _append_line(text: str):
        text = _normalize_space(text)
        if text:
            lines.append(text)

    def _para_numbering_info(p) -> Tuple[Optional[str], Optional[int]]:
        """
        Параграфтың нөмірлеу реквизиті бар-жоғын анықтау.
        Бар болса → (numId, ilvl)
        """
        try:
            pPr = p._p.pPr
            if pPr is None or pPr.numPr is None:
                return None, None
            numPr = pPr.numPr
            numId_elm = getattr(numPr, "numId", None)
            ilvl_elm = getattr(numPr, "ilvl", None)
            if numId_elm is None or ilvl_elm is None:
                return None, None
            num_id = numId_elm.val
            ilvl = _int_or_none(ilvl_elm.val)
            return str(num_id), ilvl
        except Exception:
            return None, None

    # 1) Параграфтарды оқу
    for p in doc.paragraphs:
        text = _normalize_space(p.text)
        if not text:
            continue

        num_id, ilvl = _para_numbering_info(p)
        if num_id is not None and ilvl is not None:
            marker = marker_builder.next_marker(num_id, ilvl)
            _append_line(f"{marker} {text}")
        else:
            _append_line(text)

    # 2) Кестелерді оқу (көп құжатта тесттер кестеде болады)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                # Cell ішіндегі параграфтар
                for p in cell.paragraphs:
                    text = _normalize_space(p.text)
                    if not text:
                        continue
                    num_id, ilvl = _para_numbering_info(p)
                    if num_id is not None and ilvl is not None:
                        marker = marker_builder.next_marker(num_id, ilvl)
                        _append_line(f"{marker} {text}")
                    else:
                        _append_line(text)

    return lines

# ----------------- Жылдам қолмен тексеру -----------------
if __name__ == "__main__":
    import sys, json
    if len(sys.argv) < 2:
        print("Пайдалану: python docx_parse.py <path/to/file.docx>")
        sys.exit(1)
    out = extract_text(sys.argv[1])
    print(json.dumps(out, ensure_ascii=False, indent=2))
